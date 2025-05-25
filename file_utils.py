"""
Utility functions for file processing and text extraction.
"""
import os
import re
from PyPDF2 import PdfReader
from datetime import datetime
import pandas as pd
import io

def extract_text_from_file(file_path):
    """Extract text from various file formats with improved error handling"""
    print(f"[CONTENT EXTRACTION] Starting to extract text from file: {file_path}")
    file_extension = os.path.splitext(file_path)[1].lower()
    print(f"[CONTENT EXTRACTION] File extension detected: {file_extension}")
    
    try:
        if (file_extension == '.pdf'):
            print(f"[CONTENT EXTRACTION] Processing PDF file: {file_path}")

            # Extract text from PDF with better error handling
            text = ''
            try:
                with open(file_path, 'rb') as file:
                    print(f"[CONTENT EXTRACTION] Successfully opened PDF file")
                    pdf = PdfReader(file)
                    print(f"[CONTENT EXTRACTION] PDF has {len(pdf.pages)} pages")

                    # Extract text from each page with enhanced error handling
                    for i, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if (page_text):
                                text += page_text + '\n'
                                print(f"[CONTENT EXTRACTION] Page {i+1}: Extracted {len(page_text)} characters")
                            else:
                                print(f"[CONTENT EXTRACTION] Page {i+1} returned empty text")

                                # Try alternative extraction for problematic pages
                                try:

                                    # Check if page contains image objects (scanned content)
                                    page_obj = page.get_object()
                                    if (('/Resources' in page_obj) and ('/XObject' in page_obj['/Resources'])):
                                        x_objects = page_obj['/Resources']['/XObject']
                                        if (x_objects):
                                            print(f"[CONTENT EXTRACTION] Page {i+1} contains image objects, may be scanned content")

                                            # Add placeholder for scanned content
                                            text += f"[Page {i+1} appears to contain scanned content that cannot be extracted as text]\n"
                                except Exception as obj_err:
                                    print(f"[CONTENT EXTRACTION] Error checking page objects on page {i+1}: {obj_err}")
                        except Exception as page_err:
                            print(f"[CONTENT EXTRACTION] Error extracting text from page {i+1}: {page_err}")
                            text += f"[Error extracting text from page {i+1}]\n"
            except Exception as pdf_err:
                print(f"[CONTENT EXTRACTION] Error processing PDF file: {pdf_err}")

                # Try fallback methods for problematic PDFs
                if (not text):
                    try:

                        # Another approach - try to read with a different method
                        import subprocess
                        try:

                            # Try using pdftotext if available (requires poppler)
                            print("Attempting to extract text using pdftotext")
                            result = subprocess.run(['pdftotext', file_path, '-'], 
                                                  capture_output=True, text=True, check=False)
                            if ((result.returncode == 0) and result.stdout):
                                text = result.stdout
                                print("Successfully extracted text using pdftotext")
                            else:
                                print(f"pdftotext failed with return code {result.returncode}")
                        except (FileNotFoundError, subprocess.SubprocessError) as se:
                            print(f"pdftotext not available or failed: {se}")

                            # Additional fallback: Try reading PDF in binary mode
                            try:
                                print("Attempting binary PDF read as final fallback")
                                with open(file_path, 'rb') as bin_file:
                                    raw_bytes = bin_file.read()

                                    # Extract any text-like content from binary data
                                    import re
                                    text_chunks = re.findall(b'[A-Za-z0-9 \t\r\n\f\v.,;:!?\'\"()-]{4,}', raw_bytes)
                                    if (text_chunks):
                                        decoded_chunks = [chunk.decode('utf-8', errors='replace') for chunk in text_chunks]
                                        text = '\n'.join(decoded_chunks)
                                        print(f"Extracted {len(text)} characters using binary fallback")
                            except Exception as bin_err:
                                print(f"Binary fallback also failed: {bin_err}")
                    except Exception as alt_err:
                        print(f"All alternative PDF extraction methods failed: {alt_err}")
            
            # Post-process the extracted text to clean it up
            original_length = len(text) if text else 0
            print(f"[CONTENT EXTRACTION] Raw extracted text length: {original_length} characters")
            if text:
                # Sample first 100 chars
                print(f"[CONTENT EXTRACTION] Sample of extracted text (first 100 chars): {text[:100]}")
                
                # Log occurrence of common words as content verification
                common_words = ['the', 'and', 'this', 'that', 'with', 'from']
                for word in common_words:
                    count = text.lower().count(' ' + word + ' ')
                    print(f"[CONTENT EXTRACTION] Occurrence count of '{word}': {count}")

            # Remove excessive whitespace and normalize line breaks
            import re
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\s*\n\s*', '\n', text)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            cleaned_length = len(text) if text else 0
            print(f"[CONTENT EXTRACTION] Cleaned text length: {cleaned_length} characters (removed {original_length - cleaned_length} whitespace characters)")
            
            return text
            
        elif (file_extension in ['.txt']):
            print(f"[CONTENT EXTRACTION] Processing TXT file: {file_path}")

            # Read text file with multiple encoding fallbacks
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    print(f"Failed to decode with {encoding}, trying next encoding")
                    continue
                except Exception as e:
                    print(f"Error reading text file with {encoding}: {e}")

            # If all encodings fail, try binary mode and decode with replacement
            try:
                with open(file_path, 'rb') as file:
                    binary_content = file.read()
                    return binary_content.decode('utf-8', errors='replace')
            except Exception as e:
                print(f"Binary fallback failed: {e}")
            
            return "Could not decode the text file with any supported encoding."
                
        elif (file_extension in ['.docx']):
            print(f"[CONTENT EXTRACTION] Processing DOCX file: {file_path}")

            # Use python-docx to extract text from DOCX files
            try:
                from docx import Document
                doc = Document(file_path)
                
                # Extract text from paragraphs
                full_text = []
                for para in doc.paragraphs:
                    if (para.text):

                        # Preserve paragraph formatting
                        styled_text = para.text

                        # Check for bold, italic, underline styles
                        has_formatting = False
                        for run in para.runs:
                            if (run.bold or run.italic or run.underline):
                                has_formatting = True
                                break

                        # Add a marker for paragraphs with formatting
                        if (has_formatting):
                            styled_text = f"[Formatted text] {styled_text}"
                        full_text.append(styled_text)

                # Extract text from tables
                table_count = 0
                for table in doc.tables:
                    table_count += 1
                    full_text.append(f"[Table {table_count}]")
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if (cell.text):
                                row_text.append(cell.text.strip())
                        if (row_text):
                            full_text.append(' | '.join(row_text))
                    full_text.append("[End Table]")

                # Join all text with newlines
                result = '\n'.join(full_text)

                # Check for document properties
                doc_props = []
                try:
                    core_props = doc.core_properties
                    if (core_props.title):
                        doc_props.append(f"Title: {core_props.title}")
                    if (core_props.author):
                        doc_props.append(f"Author: {core_props.author}")
                    if (core_props.last_modified_by):
                        doc_props.append(f"Last modified by: {core_props.last_modified_by}")
                    if (core_props.created):
                        doc_props.append(f"Created: {core_props.created}")
                    if (core_props.modified):
                        doc_props.append(f"Modified: {core_props.modified}")
                except:
                    pass

                # Add document properties at the beginning if available
                if (doc_props):
                    doc_props_text = "Document Properties:\n" + "\n".join(doc_props) + "\n\nDocument Content:\n"
                    result = doc_props_text + result
                
                if (not result):
                    return "The DOCX file appears to be empty or contains no extractable text."
                    
                print(f"Successfully extracted {len(result)} characters from DOCX file")
                return result
                
            except ImportError:
                print("python-docx library not available")
                return "Unable to process DOCX file: python-docx library not installed. Please install it with 'pip install python-docx'."
            except Exception as docx_err:
                print(f"Error extracting text from DOCX: {docx_err}")
                return f"Error processing DOCX file: {str(docx_err)}"
        
        elif (file_extension in ['.csv']):
            print(f"[CONTENT EXTRACTION] Processing CSV file: {file_path}")
            try:
                # Use pandas to read CSV file
                # For large CSV files, only read a sample to prevent memory issues
                # First get the file size
                file_size = os.path.getsize(file_path)
                print(f"[CONTENT EXTRACTION] CSV file size: {file_size} bytes")
                
                # If file is large (>5MB), use sampling
                MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
                SAMPLE_ROWS = 1000  # Number of rows to sample from large files
                
                if file_size > MAX_FILE_SIZE:
                    print(f"[CONTENT EXTRACTION] Large CSV detected, sampling {SAMPLE_ROWS} rows")
                    
                    # Read just the header to get column names
                    df_header = pd.read_csv(file_path, nrows=1)
                    num_columns = len(df_header.columns)
                    
                    # Read a sample of rows from the beginning
                    df_start = pd.read_csv(file_path, nrows=SAMPLE_ROWS//2)
                    
                    # Get approximate total rows to provide context
                    approx_row_size = file_size / (len(df_start) or 1) / (num_columns or 1)
                    approx_total_rows = int(file_size / approx_row_size)
                    
                    # Try to read some rows from the end (if file format allows)
                    try:
                        df_end = pd.read_csv(file_path, skiprows=range(1, approx_total_rows - SAMPLE_ROWS//2 + 1))
                        # Combine samples with a note
                        df = pd.concat([df_start, df_end])
                        note = f"\n\n[Note: This is a sample of {len(df)} rows from a large CSV with approximately {approx_total_rows:,} total rows]"
                    except Exception:
                        # If reading from the end fails, just use the beginning sample
                        df = df_start
                        note = f"\n\n[Note: This is a sample of {len(df)} rows from the beginning of a large CSV with approximately {approx_total_rows:,} total rows]"
                else:
                    # For smaller files, read the entire content
                    df = pd.read_csv(file_path)
                    note = ""
                
                # Add file metadata
                metadata = f"CSV File Analysis:\n"
                metadata += f"- Filename: {os.path.basename(file_path)}\n"
                metadata += f"- Size: {file_size:,} bytes\n"
                metadata += f"- Columns: {', '.join(df.columns.tolist())}\n"
                
                # Include basic statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    metadata += f"- Numeric column statistics:\n"
                    for col in numeric_cols[:5]:  # Limit to first 5 numeric columns
                        try:
                            metadata += f"  * {col}: min={df[col].min()}, max={df[col].max()}, mean={df[col].mean():.2f}\n"
                        except:
                            pass
                    if len(numeric_cols) > 5:
                        metadata += f"  * (statistics for {len(numeric_cols)-5} more numeric columns not shown)\n"
                
                # Convert to string representation for text extraction
                buffer = io.StringIO()
                df.to_string(buffer, index=False)
                text = metadata + "\n\nData Sample:\n" + buffer.getvalue() + note
                
                print(f"Successfully extracted {len(text)} characters from CSV file")
                return text
            except Exception as csv_err:
                print(f"Error extracting text from CSV: {csv_err}")
                return f"Error processing CSV file: {str(csv_err)}"
        
        elif (file_extension in ['.xlsx', '.xls']):
            print(f"[CONTENT EXTRACTION] Processing Excel file: {file_path}")
            try:
                # Use pandas to read Excel file
                # For large Excel files, only read a sample to prevent memory issues
                # First get the file size
                file_size = os.path.getsize(file_path)
                print(f"[CONTENT EXTRACTION] Excel file size: {file_size} bytes")
                
                # If file is large (>5MB), use sampling
                MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
                SAMPLE_ROWS = 1000  # Number of rows to sample from large files
                
                if file_size > MAX_FILE_SIZE:
                    print(f"[CONTENT EXTRACTION] Large Excel file detected, sampling {SAMPLE_ROWS} rows")
                    
                    # Get sheet names
                    xl = pd.ExcelFile(file_path)
                    sheet_names = xl.sheet_names
                    
                    # Initialize empty string for results
                    text = f"Excel File Analysis:\n"
                    text += f"- Filename: {os.path.basename(file_path)}\n"
                    text += f"- Size: {file_size:,} bytes\n"
                    text += f"- Sheets: {', '.join(sheet_names)}\n\n"
                    
                    # Process each sheet with limited rows
                    for sheet in sheet_names:
                        # Read header to get column names
                        df_header = pd.read_excel(file_path, sheet_name=sheet, nrows=1)
                        
                        # Read sample of rows
                        df = pd.read_excel(file_path, sheet_name=sheet, nrows=SAMPLE_ROWS)
                        
                        # Add sheet information
                        text += f"Sheet: {sheet} (sample of {len(df)} rows)\n"
                        text += f"Columns: {', '.join(df_header.columns.tolist())}\n\n"
                        
                        # Convert to string
                        buffer = io.StringIO()
                        df.to_string(buffer, index=False)
                        text += buffer.getvalue() + "\n\n"
                        text += f"[Note: This is a sample from sheet '{sheet}'. Full data not shown.]\n\n"
                        
                        # Set a maximum number of sheets to process to prevent excessive output
                        if len(text) > 500000:  # Limit to ~500KB of text
                            text += f"[Note: Output truncated as it exceeded size limit. Not all sheets are fully displayed.]\n"
                            break
                else:
                    # For smaller files, read the entire content with sheet names
                    xl = pd.ExcelFile(file_path)
                    sheet_names = xl.sheet_names
                    
                    text = f"Excel File Analysis:\n"
                    text += f"- Filename: {os.path.basename(file_path)}\n"
                    text += f"- Size: {file_size:,} bytes\n"
                    text += f"- Sheets: {', '.join(sheet_names)}\n\n"
                    
                    # Process each sheet (up to a reasonable limit)
                    for sheet_idx, sheet in enumerate(sheet_names):
                        if sheet_idx >= 3:  # Limit to first 3 sheets
                            text += f"[Note: {len(sheet_names) - 3} additional sheets not shown]\n"
                            break
                            
                        df = pd.read_excel(file_path, sheet_name=sheet)
                        text += f"Sheet: {sheet} ({len(df)} rows)\n"
                        
                        # Convert to string
                        buffer = io.StringIO()
                        df.head(100).to_string(buffer, index=False)  # Show only first 100 rows
                        text += buffer.getvalue()
                        
                        if len(df) > 100:
                            text += f"\n[...{len(df) - 100} more rows not shown...]\n"
                        
                        text += "\n\n"
                
                print(f"Successfully extracted {len(text)} characters from Excel file")
                return text
            except Exception as excel_err:
                print(f"Error extracting text from Excel: {excel_err}")
                return f"Error processing Excel file: {str(excel_err)}"
        
        else:
            print(f"Unsupported file extension: {file_extension}")
            return f"Unsupported file format: {file_extension}. Supported formats are: PDF, DOCX, TXT, CSV, XLSX, and XLS."
            
    except Exception as e:
        print(f"Global error in extract_text_from_file: {e}")
        return f"There was an error extracting text from this file: {str(e)}"

def read_dataframe_from_file(file_path):
    """Read data into a pandas DataFrame from CSV or Excel files"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.csv':
            return pd.read_csv(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            return pd.read_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format for data transformation: {file_extension}. Use CSV or Excel files.")
    except Exception as e:
        print(f"Error reading data into DataFrame: {e}")
        raise
